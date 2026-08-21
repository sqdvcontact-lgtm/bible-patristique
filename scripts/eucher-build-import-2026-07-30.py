from __future__ import annotations

import csv
import hashlib
import json
import random
import re
import sys
import unicodedata
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from lxml import etree


sys.stdout.reconfigure(encoding="utf-8")

MASTER = Path(r"C:\Corpus Scriptura\CS - Espace travail IA\Saint_Eucher_Du_mepris_du_monde_1672_transcription.docx")
FACSIMILE = Path(r"D:\OneDrive\Bureau\Du_mépris_du_monde.pdf")
MASTER_SHA256 = "53D61F41DD610C77875D300F81E0B50E5DE460E133AC215A49776330F706A279"
FACSIMILE_SHA256 = "4799AE77B4225144C33588FB039810EBD2412C94C9891F3E399417D0C972B261"
ROOT = Path(r"C:\Corpus Scriptura\bible-patristique\tmp\eucher-import-2026-07-30")
WORK_ID = "A0418O0003"
AUTHOR_ID = "A0418"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

SEGMENT_COLUMNS = [
    "id_oeuvre", "segment_numero", "segment_texte",
    "ref_niv1", "ref_niv2", "ref_niv3", "ref_niv4", "ref_niv5",
    "ref_niv1_texte", "ref_niv2_texte", "ref_niv3_texte", "ref_niv4_texte", "ref_niv5_texte",
    "lien_1", "lien_2", "lien_3", "lien_4", "fiabilite", "nature", "texte_original", "notes",
    "paragraphe", "rang", "controle_rang_manuel", "controle_verifie", "marquage_source",
    "liens_revus_le", "liens_revus_par",
]

ABBREVIATIONS = {"c.", "cf.", "etc.", "l.", "liv.", "m.", "mm.", "p.", "pag.", "s.", "ss.", "v."}

# Les doubles virgules des §§ 12, 15 et 16 signalent des passages remarquables,
# non des citations. On les conserve, sans leur attribuer la nature « citation ».
EDITORIAL_HIGHLIGHT_SOURCE_INDEXES = {29, 32, 33}


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest().upper()


def stable_json(value) -> str:
    return json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n"


def write_json(name: str, value):
    body = stable_json(value).encode("utf-8")
    path = ROOT / name
    path.write_bytes(body)
    (ROOT / f"{name}.sha256").write_text(f"{sha256_bytes(body)}  {name}\n", encoding="utf-8")
    return path


def xml_paragraph_text(node: etree._Element) -> str:
    out = []
    for child in node.iter():
        local = etree.QName(child).localname
        if local == "t":
            out.append(child.text or "")
        elif local == "tab":
            out.append("\t")
        elif local == "footnoteReference":
            note_id = child.get(f'{{{NS["w"]}}}id')
            out.append(f"[[WORD_FOOTNOTE_{note_id}]]")
    return "".join(out)


def normalize_typography(text: str, context: str, transformations: list[dict]) -> str:
    before = text
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\t", " ").replace("\r", " ").replace("\n", " ")
    text = re.sub(r"[ \u00a0\u202f]+", " ", text).strip()
    text = text.replace("...", "…")
    text = re.sub(r"\s+([,.])", r"\1", text)
    text = re.sub(r"\s*([:;!?])", "\u202f" + r"\1", text)
    text = re.sub(r"\s*[—–]\s*", " - ", text)
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\bS\.\s+(?=[A-ZÀ-ÖØ-Þ])", "saint ", text)
    text = re.sub(r"«\s*", "«\u202f", text)
    text = re.sub(r"\s*»", "\u202f»", text)
    if text != before:
        transformations.append({"context": context, "type": "normalisation_typographique", "before": before, "after": text})
    return text


def move_note_calls_before_punctuation(text: str, context: str, transformations: list[dict]) -> str:
    before = text
    text = re.sub(r"([,.;:!?])((?:\[\[\d+\]\])+)", r"\2\1", text)
    text = re.sub(r"([»”\"])((?:\[\[\d+\]\])+)", r"\2\1", text)
    text = re.sub(r"([,.;:!?])\s*((?:\[\[\d+\]\])+)(\u202f?»)", r"\2\1\3", text)
    if text != before:
        transformations.append({"context": context, "type": "placement_appel_note", "before": before, "after": text})
    return text


def add_quotes_once(text: str, old: str, new: str, context: str, transformations: list[dict]) -> str:
    if text.count(old) != 1:
        raise SystemExit(f"Citation introuvable ou ambiguë ({context}): {old!r}")
    updated = text.replace(old, new)
    transformations.append({"context": context, "type": "encadrement_citation", "before": old, "after": new})
    return updated


def enrich_biblical_quotes(text: str, article: int | None, transformations: list[dict]) -> str:
    replacements = {
        8: [
            ("Ces paroles de Jesus-Christ nous le font voir\u202f: Que serviroit à un homme de gagner tout le monde[[3]] & se perdre luy-mesme\u202f?",
             "Ces paroles de Jesus-Christ nous le font voir\u202f: «\u202fQue serviroit à un homme de gagner tout le monde[[3]] & se perdre luy-mesme\u202f?\u202f»"),
        ],
        19: [
            ("L’amour du bien[[5]], dit l’Apostre, est la racine de tous les maux.",
             "«\u202fL’amour du bien[[5]], dit l’Apostre, est la racine de tous les maux.\u202f»"),
        ],
        20: [
            ("Voyez ce que dit David sur ce sujet\u202f: L’avare assemble des tresors[[6]], & ne sçait pour qui il les assemble.",
             "Voyez ce que dit David sur ce sujet\u202f: «\u202fL’avare assemble des tresors[[6]], & ne sçait pour qui il les assemble.\u202f»"),
        ],
        27: [
            ("Préparons-nous à voir venir[[7]] la fin de nostre course.",
             "«\u202fPréparons-nous à voir venir[[7]] la fin de nostre course.\u202f»"),
        ],
        37: [
            ("Les ignorans ravissent le ciel[[13]]\u202f: Et nous avec toute nostre science sommes si stupides que nous demeurons toûjours ensevelis comme des bestes dans la chair & dans le sang.",
             "«\u202fLes ignorans ravissent le ciel[[13]]\u202f: Et nous avec toute nostre science sommes si stupides que nous demeurons toûjours ensevelis comme des bestes dans la chair & dans le sang.\u202f»"),
        ],
        40: [
            ("Où pourrois-je aller[[14]] pour me cacher à vostre esprit\u202f? Où pourrois-je fuir pour me dérober à vostre veuë\u202f? Si je montois dans le ciel je vous y trouverois\u202f: & si je descendois jusques dans le fond des enfers je vous y trouverois encore\u202f: si je prenois des aisles & m’envolois à l’extremité de l’orient, ou si je me retirois à l’extremité de l’occident, vous m’y viendriez prendre de vostre main pour me mener où il vous plairoit, & vous m’y tiendriez de vostre droite.",
             "«\u202fOù pourrois-je aller[[14]] pour me cacher à vostre esprit\u202f? Où pourrois-je fuir pour me dérober à vostre veuë\u202f? Si je montois dans le ciel je vous y trouverois\u202f: & si je descendois jusques dans le fond des enfers je vous y trouverois encore\u202f: si je prenois des aisles & m’envolois à l’extremité de l’orient, ou si je me retirois à l’extremité de l’occident, vous m’y viendriez prendre de vostre main pour me mener où il vous plairoit, & vous m’y tiendriez de vostre droite.\u202f»"),
        ],
        46: [
            ("Que rendray-je au Seigneur[[15]] pour tous ses bienfaits & pour toutes ses faveurs\u202f?",
             "«\u202fQue rendray-je au Seigneur[[15]] pour tous ses bienfaits & pour toutes ses faveurs\u202f?\u202f»"),
        ],
        48: [
            ("N’aimez[[16]] point le monde, dit saint Jean\u202f; ni ce qui est dans le monde",
             "«\u202fN’aimez[[16]] point le monde, dit saint Jean\u202f; ni ce qui est dans le monde\u202f»"),
            ("Les passions charnelles, comme le dit si veritablement l’Apostre saint Pierre, combattent contre l’ame[[17]]",
             "«\u202fLes passions charnelles, comme le dit si veritablement l’Apostre saint Pierre, combattent contre l’ame[[17]]\u202f»"),
        ],
        50: [
            ("Nous sommes arrivez à la fin des siecles[[18]].", "«\u202fNous sommes arrivez à la fin des siecles[[18]].\u202f»"),
        ],
        52: [
            ("Nous sommes sauvez par l’esperance[[19]].", "«\u202fNous sommes sauvez par l’esperance[[19]].\u202f»"),
        ],
        54: [
            ("Dieu l’a élevé[[20]], comme dit saint Paul, à une souveraine grandeur, & luy a donné un nom qui est au dessus de tous les noms, afin qu’au nom de Jesus tout genou fléchisse dans le ciel, sur la terre, & dans les enfers\u202f; & que toute langue confesse que le Seigneur Jesus-Christ est dans la gloire de Dieu son Pere.",
             "«\u202fDieu l’a élevé[[20]], comme dit saint Paul, à une souveraine grandeur, & luy a donné un nom qui est au dessus de tous les noms, afin qu’au nom de Jesus tout genou fléchisse dans le ciel, sur la terre, & dans les enfers\u202f; & que toute langue confesse que le Seigneur Jesus-Christ est dans la gloire de Dieu son Pere.\u202f»"),
        ],
        55: [
            ("Qu’ils n’ont des pensées[[21]] & des affections que pour la terre.",
             "«\u202fQu’ils n’ont des pensées[[21]] & des affections que pour la terre.\u202f»"),
        ],
        58: [
            ("Que l’on ne sçauroit concevoir[[22]] quels sont les biens que Dieu a préparez à ceux qui l’aiment.",
             "«\u202fQue l’on ne sçauroit concevoir[[22]] quels sont les biens que Dieu a préparez à ceux qui l’aiment.\u202f»"),
        ],
    }
    for old, new in replacements.get(article, []):
        text = add_quotes_once(text, old, new, f"article:{article}", transformations)
    return text


def word_count(text: str) -> int:
    plain = re.sub(r"\[\[\d+\]\]", "", text)
    return len(re.findall(r"[\wÀ-ÖØ-öø-ÿŒœ]+(?:[’'-][\wÀ-ÖØ-öø-ÿŒœ]+)*", plain, flags=re.UNICODE))


def quote_spans(text: str) -> list[tuple[int, int]]:
    spans = []
    stack = []
    for index, char in enumerate(text):
        if char == "«":
            stack.append(index)
        elif char == "»" and stack:
            spans.append((stack.pop(), index + 1))
    if stack:
        raise ValueError(f"Guillemet non fermé: {text}")
    return spans


def inside_quote(position: int, spans: list[tuple[int, int]]) -> bool:
    return any(start < position < end for start, end in spans)


def boundary_candidates(text: str):
    spans = quote_spans(text)
    final, strong, clause = [], [], []
    for index, char in enumerate(text):
        if char not in ".!?;:":
            continue
        end = index + 1
        trailing_note = re.match(r"(?:\[\[\d+\]\])+", text[end:])
        if trailing_note:
            end += trailing_note.end()
        if inside_quote(end, spans):
            continue
        if end < len(text) and not text[end].isspace() and text[end] != "»":
            continue
        if char == ".":
            token = re.search(r"([^\s]+)$", text[: index + 1])
            token = token.group(1).lower() if token else ""
            if token in ABBREVIATIONS or re.fullmatch(r"[ivxlcdm]+\.", token):
                continue
            if re.fullmatch(r"\d+\.", token) and int(token[:-1]) < 1000:
                continue
        (final if char in ".!?" else strong).append(end)
    for _, end in spans:
        if end == len(text) or text[end].isspace():
            final.append(end)
    clause_patterns = [
        r",\s+(?=(?:mais|car|or|donc|ny|ni|et|parce\s+que|puisque|afin\s+que|lors\s+que|lorsque|quoy\s+que|quoique)\b)",
        r",\s+(?=(?:je|tu|il|elle|nous|vous|ils|elles|on|l’on|ce|cela|c’est|c’estoit|ainsi|alors|neanmoins|toutefois|apres|avant|au\s+lieu|afin\s+de|pour|sans|en\s+[\wÀ-ÖØ-öø-ÿŒœ’'-]+ant)\b)",
        r",\s+(?=(?:luy|celuy|celle|ceux|celles|vostre|mon|ma|mes|son|sa|ses|leur|leurs)\b)",
        r",\s+(?=&\s+(?:je|tu|il|elle|nous|vous|ils|elles|on|l’on|ce|cela|ainsi|alors|en|pour|afin|sans|ne|se|s’|qu’|qui|que)\b)",
        r"\s+(?=&\s+(?:je|tu|il|elle|nous|vous|ils|elles|on|l’on|ce|cela|ainsi|alors|en|pour|afin|sans|ne|se|s’|qu’|qui|que)\b)",
        r"\s+(?=(?:qui|que|qu’|dont|où|sinon|comme|soit\s+que)\b)",
        r"\s+(?=(?:mais|car|or|donc|parce\s+que|puisque|afin\s+que|lors\s+que|lorsque)\b)",
    ]
    for pattern in clause_patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            position = match.end()
            ampersand = re.search(r"&\s*$", text[:position])
            position = ampersand.start() if ampersand else position
            if not inside_quote(position, spans):
                clause.append(position)
    return [sorted(set(final)), sorted(set(strong)), sorted(set(clause))]


def split_text(text: str, context: str, alerts: list[dict]) -> list[str]:
    candidates = boundary_candidates(text)

    def recurse(start: int, end: int, level: int = 0):
        piece = text[start:end].strip()
        count = word_count(piece)
        if count <= 18:
            return [(start, end)]
        for current_level in range(level, len(candidates)):
            usable = []
            for position in candidates[current_level]:
                if not (start < position < end):
                    continue
                left_count = word_count(text[start:position])
                right_count = word_count(text[position:end])
                min_left = 8 if current_level == 2 else 5
                min_right = 8 if current_level == 2 else 4
                if current_level == 2 and re.search(
                    r"(?:\b(?:que|qui|dont|où|mais|car|donc|de|du|des|à|au|aux|pour|sans|avec|et|ni|ny|comme|parce|puis|puisque|afin|lors|si|ce|cela|quoy|quoi|ou|soit|bien|encore|plus|moins|tant|autant|tel|telle|tels|telles|mesme|même|pourveu|ainsi)|&)\s*[,;:]?$",
                    text[start:position].rstrip(), flags=re.IGNORECASE,
                ):
                    continue
                next_text = text[position:end].lstrip()
                if current_level == 2 and re.match(r"(?:que|qu’)\b", next_text, flags=re.IGNORECASE) and re.search(
                    r"(?:\bne\b|\bn’)[^,;:.!?]{0,90}$", text[start:position].rstrip(), flags=re.IGNORECASE,
                ):
                    continue
                if left_count >= min_left and right_count >= min_right:
                    usable.append((position, left_count, right_count))
            if not usable:
                continue
            position, _, _ = min(usable, key=lambda item: (abs(item[1] - 14), max(item[1], item[2]), item[0]))
            return recurse(start, position, current_level) + recurse(position, end, current_level)
        if count > 35:
            alerts.append({"type": "segment_long_conserve", "context": context, "words": count, "text": piece})
        return [(start, end)]

    forced_boundaries = []
    for _, quote_end in quote_spans(text):
        position = quote_end
        while position < len(text) and text[position] in ",;:.!?":
            position += 1
        if position == len(text) or text[position].isspace():
            forced_boundaries.append(position)
    spans = []
    forced_start = 0
    for forced_end in sorted(set(forced_boundaries)):
        if forced_start < forced_end:
            spans.extend(recurse(forced_start, forced_end))
        forced_start = forced_end
    if forced_start < len(text):
        spans.extend(recurse(forced_start, len(text)))
    pieces = [text[start:end].strip() for start, end in spans if text[start:end].strip()]
    if " ".join(pieces) != text:
        raise ValueError(f"Recomposition brisée: {context}")
    return pieces


master_bytes = MASTER.read_bytes()
facsimile_bytes = FACSIMILE.read_bytes()
if sha256_bytes(master_bytes) != MASTER_SHA256:
    raise SystemExit("Le Word maître a changé depuis sa validation")
if sha256_bytes(facsimile_bytes) != FACSIMILE_SHA256:
    raise SystemExit("Le fac-similé a changé depuis son identification")
ROOT.mkdir(parents=True, exist_ok=True)

doc = Document(MASTER)
with zipfile.ZipFile(MASTER) as archive:
    document_root = etree.fromstring(archive.read("word/document.xml"))
    footnotes_root = etree.fromstring(archive.read("word/footnotes.xml"))
xml_paragraphs = document_root.xpath("//w:body/w:p", namespaces=NS)
if len(xml_paragraphs) != len(doc.paragraphs):
    raise SystemExit("Divergence entre les paragraphes DOCX et OOXML")

footnotes_by_id = {}
for node in footnotes_root.xpath("//w:footnote", namespaces=NS):
    note_id = node.get(f'{{{NS["w"]}}}id')
    if note_id and int(note_id) > 0:
        footnotes_by_id[note_id] = "".join(node.xpath(".//w:t/text()", namespaces=NS)).strip()
if len(footnotes_by_id) != 22:
    raise SystemExit(f"22 notes Word attendues, {len(footnotes_by_id)} trouvées")

transformations = []
alerts = []
paragraphs = []
ordered_internal_note_ids = []
for index, (paragraph, xml_node) in enumerate(zip(doc.paragraphs, xml_paragraphs)):
    xml_text = xml_paragraph_text(xml_node)
    if re.sub(r"\[\[WORD_FOOTNOTE_\d+\]\]", "", xml_text) != paragraph.text:
        raise SystemExit(f"Divergence texte/OOXML au paragraphe {index}")
    ordered_internal_note_ids.extend(re.findall(r"\[\[WORD_FOOTNOTE_(\d+)\]\]", xml_text))
    paragraphs.append({"index": index, "style": paragraph.style.name, "raw_text": xml_text})
if len(ordered_internal_note_ids) != 22 or len(set(ordered_internal_note_ids)) != 22:
    raise SystemExit(f"Ordre des appels Word inattendu: {ordered_internal_note_ids}")

note_number_by_internal_id = {internal_id: number for number, internal_id in enumerate(ordered_internal_note_ids, start=1)}
note_definitions = {
    note_number_by_internal_id[internal_id]: normalize_typography(text, f"note:{internal_id}", transformations)
    for internal_id, text in footnotes_by_id.items()
}
for paragraph in paragraphs:
    text = paragraph["raw_text"]
    for internal_id, number in note_number_by_internal_id.items():
        text = text.replace(f"[[WORD_FOOTNOTE_{internal_id}]]", f"[[{number}]]")
    paragraph["clean_text"] = move_note_calls_before_punctuation(
        normalize_typography(text, f"docx:{paragraph['index']}", transformations),
        f"docx:{paragraph['index']}",
        transformations,
    ) if text else ""
if any("WORD_FOOTNOTE" in paragraph["clean_text"] for paragraph in paragraphs):
    raise SystemExit("Appel Word résiduel")

segments = []
source_map = []
segment_number = 0


def push_paragraph(text: str, nature: str, ref1: str | None, ref1_text: str | None,
                   paragraph_number: int, source_index: int, article_source: int | None = None):
    global segment_number
    pieces = split_text(text, f"docx:{source_index}", alerts)
    first_segment = segment_number + 1
    for rank, piece in enumerate(pieces, start=1):
        segment_number += 1
        calls = [int(value) for value in re.findall(r"\[\[(\d+)\]\]", piece)]
        definitions = "\n".join(f"[[{number}]] {note_definitions[number]}" for number in calls) or None
        segment_nature = nature
        if nature == "texte" and source_index not in EDITORIAL_HIGHLIGHT_SOURCE_INDEXES and "«" in piece and "»" in piece:
            segment_nature = "citation"
        segments.append({
            "id_oeuvre": WORK_ID,
            "segment_numero": segment_number,
            "segment_texte": piece,
            "ref_niv1": ref1,
            "ref_niv2": None,
            "ref_niv3": None,
            "ref_niv4": None,
            "ref_niv5": None,
            "ref_niv1_texte": ref1_text,
            "ref_niv2_texte": None,
            "ref_niv3_texte": None,
            "ref_niv4_texte": None,
            "ref_niv5_texte": None,
            "lien_1": None,
            "lien_2": None,
            "lien_3": None,
            "lien_4": None,
            "fiabilite": None,
            "nature": segment_nature,
            "texte_original": None,
            "notes": definitions,
            "paragraphe": paragraph_number,
            "rang": rank,
            "controle_rang_manuel": None,
            "controle_verifie": False,
            "marquage_source": "Codex (IA)",
            "liens_revus_le": None,
            "liens_revus_par": None,
        })
    source_map.append({
        "source_index": source_index,
        "article_source": article_source,
        "nature": nature,
        "ref_niv1": ref1,
        "ref_niv1_texte": ref1_text,
        "paragraphe": paragraph_number,
        "source_clean": text,
        "first_segment_numero": first_segment,
        "last_segment_numero": segment_number,
        "segment_count": len(pieces),
    })


apparatus_ranges = [
    ("Abrégé de la vie de saint Eucher", "Tiré du Martyrologe d’Adon", range(7, 11)),
    ("Avertissement", None, range(12, 17)),
    ("Approbation des docteurs", None, range(81, 86)),
    ("Extrait du privilège du Roy", None, range(87, 92)),
]
for ref1, ref1_text, indexes in apparatus_ranges:
    paragraph_number = 0
    for index in indexes:
        text = paragraphs[index]["clean_text"]
        if not text:
            continue
        paragraph_number += 1
        push_paragraph(text, "apparat_critique", ref1, ref1_text, paragraph_number, index)

body_paragraph = 0
article_paragraph = 0
last_article = 0
for index in range(18, 79):
    text = paragraphs[index]["clean_text"]
    match = re.match(r"^(\d+)\.\s+(.*)$", text)
    article = int(match.group(1)) if match else None
    if article is not None:
        if article != last_article + 1:
            raise SystemExit(f"Suite des articles interrompue au DOCX {index}: {article} après {last_article}")
        last_article = article
        article_paragraph = 1
        text = match.group(2)
    elif index != 75 or last_article != 57:
        raise SystemExit(f"Paragraphe non numéroté inattendu dans le corps: {index}")
    else:
        article_paragraph += 1
    body_paragraph += 1
    if index in EDITORIAL_HIGHLIGHT_SOURCE_INDEXES:
        if text.count("«") != 1 or text.count("»") != 1:
            raise SystemExit(f"Doubles virgules marginales inattendues au DOCX {index}")
        before_highlight = text
        text = text.replace("«\u202f", "", 1).replace("\u202f»", "", 1)
        transformations.append({
            "context": f"docx:{index}",
            "type": "retrait_doubles_virgules_marginales_non_textuelles",
            "before": before_highlight,
            "after": text,
            "reason": "marques marginales de passage remarquable, non guillemets de citation",
        })
    text = enrich_biblical_quotes(text, article, transformations)
    push_paragraph(text, "texte", str(last_article), None, article_paragraph, index, article)
if last_article != 60 or body_paragraph != 61:
    raise SystemExit(f"Corps incomplet: article final {last_article}, paragraphes {body_paragraph}")

if paragraphs[79]["clean_text"] != "Fin.":
    raise SystemExit("Marque de fin inattendue")
printer_colophon = paragraphs[92]["clean_text"]
expected_colophon = "Achevé d’imprimer pour la premiere fois le troisiéme Decembre 1671."
if printer_colophon != expected_colophon:
    raise SystemExit(f"Colophon inattendu: {printer_colophon}")

calls = [int(value) for row in segments for value in re.findall(r"\[\[(\d+)\]\]", row["segment_texte"])]
definitions = [int(value) for row in segments for value in re.findall(r"^\[\[(\d+)\]\]", row["notes"] or "", flags=re.MULTILINE)]
if calls != list(range(1, 23)) or definitions != list(range(1, 23)):
    raise SystemExit(f"Séquence des notes invalide: appels={calls}, définitions={definitions}")

segment_by_number = {row["segment_numero"]: row for row in segments}
recomposition_errors = []
rank_errors = []
for mapping in source_map:
    rows = [segment_by_number[number] for number in range(mapping["first_segment_numero"], mapping["last_segment_numero"] + 1)]
    recomposed = " ".join(row["segment_texte"] for row in rows)
    if recomposed != mapping["source_clean"]:
        recomposition_errors.append({"source_index": mapping["source_index"], "expected": mapping["source_clean"], "actual": recomposed})
    if [row["rang"] for row in rows] != list(range(1, len(rows) + 1)):
        rank_errors.append(mapping["source_index"])

grouped = defaultdict(list)
for row in segments:
    grouped[(row["nature"] == "apparat_critique", row["ref_niv1"], row["paragraphe"])].append(row)
group_rank_errors = []
for key, rows in grouped.items():
    ranks = sorted(row["rang"] for row in rows)
    if ranks != list(range(1, len(rows) + 1)):
        group_rank_errors.append({"key": key, "ranks": ranks})

rng = random.Random(20260730)
random_checks = []
for mapping in rng.sample(source_map, 14):
    rows = [segment_by_number[number] for number in range(mapping["first_segment_numero"], mapping["last_segment_numero"] + 1)]
    random_checks.append({
        "source_index": mapping["source_index"],
        "article_source": mapping["article_source"],
        "nature": mapping["nature"],
        "ref_niv1": mapping["ref_niv1"],
        "paragraphe": mapping["paragraphe"],
        "segments": [row["segment_numero"] for row in rows],
        "recomposition_ok": " ".join(row["segment_texte"] for row in rows) == mapping["source_clean"],
        "start": mapping["source_clean"][:180],
    })

nb_signes = sum(len(row["segment_texte"]) for row in segments)
final_note = (
    "Traduction française de Robert Arnauld d’Andilly, Paris, Pierre Le Petit, 1672. "
    "L’édition imprime le texte latin à la suite de la traduction. Mention d’achevé d’imprimer "
    f"conservée dans la notice bibliographique : « {printer_colophon} »"
)
metadata = {
    "oeuvre_initiale": {
        "id_oeuvre": WORK_ID,
        "id_auteur": AUTHOR_ID,
        "titre": "Du mépris du monde",
        "sous_titre": "Lettre de saint Eucher, Evesque de Lyon, à Valere",
        "titre_original": "De contemptu mundi",
        "langue_originale": "Latin",
        "langue_trad": "Français",
        "date_approx": "Ve siècle",
        "genre": "Lettre parénétique",
        "trad_auteur": "Robert Arnauld d’Andilly",
        "note": "[Corpus Scriptura:depublie]",
        "editeur": "Pierre Le Petit",
        "collection": None,
        "ville": "Paris",
        "trad_id": None,
        "date_publication": "1672",
        "url_source": None,
        "profondeur_sommaire": 1,
        "nb_signes": nb_signes,
        "niveaux_sommaire": 1,
        "niveaux_corps": 1,
        "texte_sommaire": "1,0,0,0,0",
        "texte_corps": "1,0,0,0,0",
        "afficher_numeros": False,
        "date_composition": "Ve siècle",
        "genres": ["Lettre", "Ascétique"],
        "composition_debut_annee": None,
        "composition_debut_precision": None,
        "composition_fin_annee": None,
        "composition_fin_precision": None,
        "publication_debut_annee": 1672,
        "publication_debut_precision": "exacte",
        "publication_fin_annee": 1672,
        "publication_fin_precision": "exacte",
        "titre_affichage": "Du mépris\ndu monde",
        "date_mise_en_ligne": None,
        "commentaire_traduction": None,
    },
    "oeuvre_apres_publication": {"note": final_note},
    "catalogue_notice": {
        "id_auteur": AUTHOR_ID,
        "auteur": "Eucher de Lyon",
        "dates_auteur": "Vers 370-Vers 450",
        "id_oeuvre_stable": WORK_ID,
        "titre_stable": "Du mépris du monde",
        "titre_original": "De contemptu mundi",
        "genre": "lettre parénétique et ascétique",
        "langue_originale": "latin",
        "date_oeuvre": "Ve siècle",
        "authenticite": "authentique",
        "id_traduction": "TR_FR_1672_ARNAULD_DANDILLY_DE_CONTEMPTU_MUNDI",
        "titre_edition": "S. Eucher du mépris du monde",
        "traducteur": "Robert Arnauld d’Andilly",
        "annee_edition": 1672,
        "siecle_edition": "XVIIe siècle",
        "editeur": "Pierre Le Petit",
        "collection_nom": None,
        "domaine_public": "oui",
        "url_source": None,
        "decision_import": "Import préparé - texte français complet, structuré et contrôlé ; liens bibliques différés.",
        "niveau_verification": (
            f"Très fort - fac-similé local SHA-256 {FACSIMILE_SHA256} ; Word maître SHA-256 {MASTER_SHA256}. "
            f"Structure, 61 paragraphes du corps, 22 notes et recomposition contrôlés. {printer_colophon}"
        ),
        "score_fiabilite": 95,
        "presence_sur_le_site": False,
        "priorite": "Haute",
        "verifie": False,
        "lieu_edition": "Paris",
        "date_edition": "1672",
        "auteur_debut_annee": 370,
        "auteur_debut_precision": "vers",
        "auteur_fin_annee": 450,
        "auteur_fin_precision": "vers",
        "oeuvre_debut_annee": None,
        "oeuvre_debut_precision": None,
        "oeuvre_fin_annee": None,
        "oeuvre_fin_precision": None,
        "edition_debut_annee": 1672,
        "edition_debut_precision": "exacte",
        "edition_fin_annee": 1672,
        "edition_fin_precision": "exacte",
        "verifie_admin": False,
        "auteur_uniformise": "Eucher de Lyon",
        "traducteur_uniformise": "Robert Arnauld d’Andilly",
        "pseudonyme_ou_nom_imprime_traducteur": "Arnauld d’Andilly",
        "refuse_admin": False,
        "url_texte_integral": None,
    },
    "bibliographic_colophon": printer_colophon,
}

page_title_fragments = [paragraphs[index]["clean_text"] for index in range(0, 5)]
all_segment_text = "\n".join(row["segment_texte"] for row in segments)
invariants = {
    "word_hash_matches": sha256_bytes(master_bytes) == MASTER_SHA256,
    "facsimile_hash_matches": sha256_bytes(facsimile_bytes) == FACSIMILE_SHA256,
    "page_title_excluded": all(fragment not in all_segment_text for fragment in page_title_fragments if fragment),
    "work_title_not_duplicated_as_level": all(row["ref_niv1"] != "Du mépris du monde" for row in segments),
    "body_articles_are_niv1": sorted(
        {int(row["ref_niv1"]) for row in segments if row["nature"] in ("texte", "citation")}
    ) == list(range(1, 61)),
    "article_57_has_two_source_paragraphs": sorted({
        row["paragraphe"] for row in segments
        if row["nature"] in ("texte", "citation") and row["ref_niv1"] == "57"
    }) == [1, 2],
    "other_articles_have_one_source_paragraph": all(
        {row["paragraphe"] for row in segments if row["nature"] in ("texte", "citation") and row["ref_niv1"] == str(article)} == {1}
        for article in range(1, 61) if article != 57
    ),
    "articles_1_to_60_complete": last_article == 60,
    "body_source_paragraphs_61": body_paragraph == 61,
    "notes_1_to_22_unique": calls == list(range(1, 23)) and definitions == list(range(1, 23)),
    "note_calls_before_punctuation": all(not re.search(r"[,.;:!?»]\[\[\d+\]\]", row["segment_texte"]) for row in segments),
    "notes_never_after_closing_quote": all(not re.search(r"[»”\"]\s*\[\[\d+\]\]", row["segment_texte"]) for row in segments),
    "segment_numbers_contiguous": [row["segment_numero"] for row in segments] == list(range(1, len(segments) + 1)),
    "paragraph_recomposition_exact": not recomposition_errors,
    "ranks_contiguous": not rank_errors and not group_rank_errors,
    "no_empty_text": all(row["segment_texte"].strip() for row in segments),
    "no_links_in_legacy_columns": all(not any(row[column] for column in ("lien_1", "lien_2", "lien_3", "lien_4")) for row in segments),
    "links_not_reviewed": all(row["liens_revus_le"] is None and row["liens_revus_par"] is None for row in segments),
    "colophon_excluded_from_segments": printer_colophon not in all_segment_text,
    "all_random_recompositions_pass": all(item["recomposition_ok"] for item in random_checks),
}
if not all(invariants.values()):
    raise SystemExit(f"Audit candidat en échec: {json.dumps(invariants, ensure_ascii=False)}")

audit = {
    "work_id": WORK_ID,
    "source": {
        "docx": str(MASTER),
        "docx_sha256": MASTER_SHA256,
        "facsimile": str(FACSIMILE),
        "facsimile_sha256": FACSIMILE_SHA256,
        "facsimile_pages": 138,
        "french_body_pdf_pages": "19-78 environ ; ouverture vérifiée aux pages PDF 17-21",
    },
    "structure": {
        "excluded_title_page_paragraphs": list(range(0, 5)),
        "apparatus_sections": [item[0] for item in apparatus_ranges],
        "body_articles": 60,
        "body_source_paragraphs": 61,
        "apparatus_paragraphs": sum(1 for item in source_map if item["nature"] == "apparat_critique"),
        "segments": len(segments),
        "body_segments": sum(row["nature"] in ("texte", "citation") for row in segments),
        "apparatus_segments": sum(row["nature"] == "apparat_critique" for row in segments),
        "citation_segments": sum(row["nature"] == "citation" for row in segments),
        "notes": 22,
        "nb_signes": nb_signes,
        "nature_counts": Counter(row["nature"] for row in segments),
    },
    "invariants": invariants,
    "recomposition_errors": recomposition_errors,
    "rank_errors": rank_errors,
    "group_rank_errors": group_rank_errors,
    "alerts": alerts,
    "random_seed": 20260730,
    "random_checks": random_checks,
    "stop_before_links": True,
}

write_json("eucher-segments-candidate.json", segments)
write_json("eucher-source-map.json", source_map)
write_json("eucher-metadata-candidate.json", metadata)
write_json("eucher-transformations.json", transformations)
write_json("eucher-alerts.json", alerts)
write_json("eucher-random-checks.json", random_checks)
write_json("eucher-audit.json", audit)

csv_path = ROOT / "eucher-segments-candidate.csv"
with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
    writer = csv.DictWriter(handle, fieldnames=["id"] + SEGMENT_COLUMNS)
    writer.writeheader()
    for row in segments:
        writer.writerow({"id": "", **row})

print(stable_json({
    "ready": True,
    "work_id": WORK_ID,
    "segments": len(segments),
    "body_paragraphs": body_paragraph,
    "apparatus_paragraphs": audit["structure"]["apparatus_paragraphs"],
    "notes": len(note_definitions),
    "alerts": len(alerts),
    "invariants": invariants,
    "candidate_sha256": sha256_bytes((ROOT / "eucher-segments-candidate.json").read_bytes()),
}))
