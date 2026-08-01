from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from docx import Document


MASTER = Path(r"C:\Corpus Scriptura\CS - Espace travail IA\Saint_Eucher_Du_mepris_du_monde_1672_transcription.docx")
BACKUP = MASTER.with_name(f"{MASTER.stem}_avant_correction_audit_2026-07-30.docx")
EXPECTED_BEFORE_SHA256 = "59AECCA3C1FAE633FB87BF51A5F6B27FEAE760F6343EEEF3700574C4F1DC3F27"

REPLACEMENTS = {
    "A. Darrera Curé de S. André.": "A. Debreda Curé de S. André.",
    "Gremet Curé de S. Benoist.": "Grenet Curé de S. Benoist.",
    "N. Gorillon Curé de S. Laurent.": "N. Gobillon Curé de S. Laurent.",
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


before_hash = sha256(MASTER)
if before_hash != EXPECTED_BEFORE_SHA256:
    raise SystemExit(f"Word maître inattendu : {before_hash}")

if BACKUP.exists():
    if sha256(BACKUP) != EXPECTED_BEFORE_SHA256:
        raise SystemExit(f"Sauvegarde existante inattendue : {BACKUP}")
else:
    shutil.copy2(MASTER, BACKUP)

document = Document(MASTER)
counts = {old: 0 for old in REPLACEMENTS}
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        for old, new in REPLACEMENTS.items():
            found = run.text.count(old)
            if found:
                run.text = run.text.replace(old, new)
                counts[old] += found

if any(count != 1 for count in counts.values()):
    raise SystemExit(f"Occurrences inattendues : {counts}")

document.save(MASTER)
reopened = Document(MASTER)
corpus = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
for old, new in REPLACEMENTS.items():
    if old in corpus or corpus.count(new) != 1:
        raise SystemExit(f"Validation échouée : {old!r} -> {new!r}")

print(f"backup={BACKUP}")
print(f"before_sha256={before_hash}")
print(f"after_sha256={sha256(MASTER)}")
print(f"replacements={counts}")
