from __future__ import annotations

import hashlib
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


SOURCE = Path(r"C:\Corpus Scriptura\Sources\Basile\BKV_Homelies_sur_Hexaemeron_Auger_1827.docx")
OUTPUT = Path(r"C:\Corpus Scriptura\bible-patristique\audit\hexameron-2026-07-30\bkv-docx-structure.json")
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def paragraph_text(paragraph) -> str:
    return re.sub(r"\s+", " ", paragraph.text).strip()


def footnotes(path: Path) -> list[dict[str, object]]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/footnotes.xml"))
    notes: list[dict[str, object]] = []
    for note in root.findall(f"{W}footnote"):
        note_id = note.get(f"{W}id")
        if note_id in {"-1", "0"}:
            continue
        text = "".join(node.text or "" for node in note.iter(f"{W}t"))
        notes.append({"id": int(note_id), "text": re.sub(r"\s+", " ", text).strip()})
    return notes


def main() -> None:
    document = Document(SOURCE)
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph_text(paragraph)
        if not text:
            continue
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else None,
                "text": text,
            }
        )

    result = {
        "source": str(SOURCE),
        "sha256": hashlib.sha256(SOURCE.read_bytes()).hexdigest().upper(),
        "paragraph_count_all": len(document.paragraphs),
        "paragraph_count_nonempty": len(paragraphs),
        "style_counts_nonempty": dict(Counter(item["style"] for item in paragraphs)),
        "paragraphs": paragraphs,
        "footnotes": footnotes(SOURCE),
    }
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({key: value for key, value in result.items() if key not in {"paragraphs", "footnotes"}}, ensure_ascii=False, indent=2))
    print(f"footnotes={len(result['footnotes'])}")
    print("first paragraphs:")
    for item in paragraphs[:30]:
        print(f"{item['index']:4} | {item['style']!s:24} | {item['text'][:180]}")


if __name__ == "__main__":
    main()
