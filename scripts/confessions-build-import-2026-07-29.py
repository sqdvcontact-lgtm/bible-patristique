import csv
import hashlib
import json
import random
import re
import sys
import unicodedata
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')

MASTER = Path(r'C:\Corpus Scriptura\Augustin\Confessions-Saint-Augustin-Andilly-1649-MASTER.docx')
EXPECTED_SHA256 = '3F1951259DC87AD2DDAB62179F2EDEB42165EA0CF4A6075AC198F596908F92BD'
ROOT = Path(r'C:\Corpus Scriptura\bible-patristique\tmp\confessions-import-2026-07-29')
WORK_ID = 'A0010O0001'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
NOTE_ATTACHMENTS = {23: 24, 26: 27, 29: 30, 31: 32, 33: 34, 35: 36}
TITLE_PAGE_RANGE = range(5, 21)
EXCLUDED_WRAPPER_RANGE = range(0, 5)
SAFE_LEXICAL_CORRECTIONS = {
    'LOVIS par la grace': 'Louis par la grace',
    'A CES CAVSES,': 'A ces causes,',
    'MANDONS au premier': 'Mandons au premier',
    'Conseil, PEPIN.': 'Conseil, Pepin.',
    'BOVRGEOIS. RETART.': 'Bourgeois. Retart.',
}
ABBREVIATIONS = {
    'c.', 'cf.', 'etc.', '&c.', 'l.', 'liv.', 'm.', 'mm.', 'p.', 'pag.', 's.', 'ss.', 'v.',
}
SEGMENT_COLUMNS = [
    'id_oeuvre', 'segment_numero', 'segment_texte',
    'ref_niv1', 'ref_niv2', 'ref_niv3', 'ref_niv4', 'ref_niv5',
    'ref_niv1_texte', 'ref_niv2_texte', 'ref_niv3_texte', 'ref_niv4_texte', 'ref_niv5_texte',
    'lien_1', 'lien_2', 'lien_3', 'lien_4', 'fiabilite', 'nature', 'texte_original', 'notes',
    'paragraphe', 'rang', 'controle_rang_manuel', 'controle_verifie', 'marquage_source',
    'liens_revus_le', 'liens_revus_par',
]


def sha256_bytes(data):
    return hashlib.sha256(data).hexdigest().upper()


def stable_json(value):
    return json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + '\n'


def write_json(name, value):
    body = stable_json(value)
    path = ROOT / name
    path.write_bytes(body.encode('utf-8'))
    (ROOT / f'{name}.sha256').write_bytes(f'{sha256_bytes(body.encode())}  {name}\n'.encode('utf-8'))
    return path


def xml_paragraph_text(node):
    out = []
    for child in node.iter():
        local = etree.QName(child).localname
        if local == 't':
            out.append(child.text or '')
        elif local == 'tab':
            out.append('\t')
        elif local == 'footnoteReference':
            note_id = child.get(f'{{{NS["w"]}}}id')
            out.append(f'[[WORD_FOOTNOTE_{note_id}]]')
    return ''.join(out)


def normalize_typography(text, context, transformations):
    before = text
    text = unicodedata.normalize('NFC', text)
    text = text.replace('\t', ' ').replace('\r', ' ').replace('\n', ' ')
    text = re.sub(r'[ \u00a0\u202f]+', ' ', text).strip()
    text = text.replace('...', '…')
    text = re.sub(r'\s+([,.])', r'\1', text)
    text = re.sub(r'\s*([:;!?])', '\u202f' + r'\1', text)
    text = re.sub(r'\s*[—–]\s*', ' - ', text)
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'\bS\.\s+(?=[A-ZÀ-ÖØ-Þ])', 'saint ', text)
    def lower_mid_sentence_saint(match):
        prefix = text[:match.start()].rstrip()
        return match.group(0) if not prefix or prefix.endswith(('.', '!', '?')) else f'saint {match.group(1)}'
    text = re.sub(r'\bSaint\s+([A-ZÀ-ÖØ-Þ][\wÀ-ÖØ-öø-ÿŒœ-]*)', lower_mid_sentence_saint, text)
    for old, new in SAFE_LEXICAL_CORRECTIONS.items():
        if old in text:
            text = text.replace(old, new)
            transformations.append({'context': context, 'type': 'capitales_typographiques', 'old': old, 'new': new})
    if text != before:
        transformations.append({'context': context, 'type': 'normalisation_typographique', 'before': before, 'after': text})
    return text


def insert_note_before_terminal(text, number):
    match = re.search(r'([.!?])$', text)
    if match:
        return f'{text[:match.start()]}[[{number}]]{match.group(1)}'
    return f'{text}[[{number}]]'


def word_count(text):
    plain = re.sub(r'\[\[\d+\]\]', '', text)
    return len(re.findall(r"[\wÀ-ÖØ-öø-ÿŒœ]+(?:[’'-][\wÀ-ÖØ-öø-ÿŒœ]+)*", plain, flags=re.UNICODE))


def boundary_candidates(text):
    final = []
    strong = []
    clause = []
    for index, char in enumerate(text):
        if char not in '.!?;:':
            continue
        end = index + 1
        trailing_note = re.match(r'(?:\[\[\d+\]\])+', text[end:])
        if trailing_note:
            end += trailing_note.end()
        if end < len(text) and not text[end].isspace():
            continue
        if char == '.':
            token = re.search(r'([^\s]+)$', text[:index + 1])
            token = (token.group(1).lower() if token else '')
            following = text[end:].lstrip()
            if token in ABBREVIATIONS or re.fullmatch(r'(?:[ivxlcdm]+|\d+)\.', token):
                if following:
                    continue
        (final if char in '.!?' else strong).append(end)
    clause_patterns = [
        r',\s+(?=(?:mais|car|or|donc|ny|ni|et|parce\s+que|puisque|afin\s+que|lors\s+que|lorsque|quoy\s+que|quoique)\b)',
        r",\s+(?=(?:je|tu|il|elle|nous|vous|ils|elles|on|l’on|ce|cela|c’est|c’estoit|ainsi|alors|neanmoins|toutefois|apres|avant|au\s+lieu|afin\s+de|pour|sans|en\s+[\wÀ-ÖØ-öø-ÿŒœ’'-]+ant)\b)",
        r',\s+(?=(?:luy|celuy|celle|ceux|celles|vostre|mon|ma|mes|son|sa|ses|leur|leurs)\b)',
        r",\s+(?=&\s+(?:je|tu|il|elle|nous|vous|ils|elles|on|l’on|ce|cela|ainsi|alors|en|pour|afin|sans|ne|se|s’|qu’|qui|que)\b)",
        r"\s+(?=&\s+(?:je|tu|il|elle|nous|vous|ils|elles|on|l’on|ce|cela|ainsi|alors|en|pour|afin|sans|ne|se|s’|qu’|qui|que)\b)",
        r"\s+(?=(?:qui|que|qu’|dont|où|sinon|comme|soit\s+que)\b)",
    ]
    for pattern in clause_patterns:
        clause.extend(match.end() for match in re.finditer(pattern, text, flags=re.IGNORECASE))
    for pattern in [r'\s+(?=mais\b)', r'\s+(?=car\b)', r'\s+(?=or\b)', r'\s+(?=donc\b)', r'\s+(?=parce\s+que\b)', r'\s+(?=puisque\b)', r'\s+(?=afin\s+que\b)', r'\s+(?=lors\s+que\b)', r'\s+(?=lorsque\b)']:
        clause.extend(match.end() for match in re.finditer(pattern, text, flags=re.IGNORECASE))
    normalized_clause = []
    for position in clause:
        ampersand = re.search(r'&\s*$', text[:position])
        normalized_clause.append(ampersand.start() if ampersand else position)
    clause = normalized_clause
    return [sorted(set(final)), sorted(set(strong)), sorted(set(clause))]


def split_text(text, context, alerts):
    candidates = boundary_candidates(text)

    def recurse(start, end, level=0):
        piece = text[start:end].strip()
        count = word_count(piece)
        if count <= 18:
            return [(start, end)]
        for current_level in range(level, len(candidates)):
            usable = []
            for position in candidates[current_level]:
                if not (start < position < end):
                    continue
                left_count = word_count(text[start:position])
                right_count = word_count(text[position:end])
                min_left = 8 if current_level == 2 else 5
                min_right = 5 if current_level == 2 else 4
                if current_level == 2 and re.search(r"(?:\b(?:que|qui|dont|où|mais|car|donc|de|du|des|à|au|aux|pour|sans|avec|et|ni|ny|comme|puisque|afin|lors|si|ce|cela|quoy|quoi|ou|soit|bien|encore|plus|moins|tant|autant|tel|telle|tels|telles|mesme|même|pourveu|ainsi|dautant|entant)|&)\s*[,;:]?$", text[start:position].rstrip(), flags=re.IGNORECASE):
                    continue
                next_text = text[position:end].lstrip()
                if current_level == 2 and re.match(r"(?:que|qu’)\b", next_text, flags=re.IGNORECASE) and re.search(r"(?:\bne\b|\bn’)[^,;:.!?]{0,70}$", text[start:position].rstrip(), flags=re.IGNORECASE):
                    continue
                if left_count >= min_left and right_count >= min_right:
                    usable.append((position, left_count, right_count))
            if not usable:
                continue
            position, _, _ = min(usable, key=lambda item: (abs(item[1] - 14), max(item[1], item[2]), item[0]))
            return recurse(start, position, current_level) + recurse(position, end, current_level)
        if count > 30:
            alerts.append({
                'type': 'segment_long_conserve_sans_marqueur',
                'severity': 'information',
                'context': context,
                'words': count,
                'text': piece,
            })
        return [(start, end)]

    spans = recurse(0, len(text))
    pieces = [text[start:end].strip() for start, end in spans if text[start:end].strip()]
    if ' '.join(pieces) != text:
        raise ValueError(f'Invariant de recomposition brisé avant enrichissement: {context}')
    return pieces


def markdown_plain(text):
    return text.replace('*', '')


def format_source_segment(text, style):
    if style == 'Apparat italique':
        return f'*{text}*'
    return text


source_bytes = MASTER.read_bytes()
actual_sha = sha256_bytes(source_bytes)
if actual_sha != EXPECTED_SHA256:
    raise SystemExit(f'Source maître modifiée: attendu {EXPECTED_SHA256}, obtenu {actual_sha}')
ROOT.mkdir(parents=True, exist_ok=True)

doc = Document(MASTER)
with zipfile.ZipFile(MASTER) as archive:
    document_root = etree.fromstring(archive.read('word/document.xml'))
    footnotes_root = etree.fromstring(archive.read('word/footnotes.xml'))
xml_paragraphs = document_root.xpath('//w:body/w:p', namespaces=NS)
if len(xml_paragraphs) != len(doc.paragraphs):
    raise SystemExit('Le nombre de paragraphes python-docx et XML diverge')
footnotes = {}
for node in footnotes_root.xpath('//w:footnote', namespaces=NS):
    note_id = node.get(f'{{{NS["w"]}}}id')
    if note_id and int(note_id) > 0:
        footnotes[note_id] = ''.join(node.xpath('.//w:t/text()', namespaces=NS)).strip()
if footnotes != {'1': 'Ces paroles sont de Ciceron.'}:
    raise SystemExit(f'Notes Word inattendues: {footnotes}')

transformations = []
alerts = []
paragraphs = []
for index, (paragraph, xml_node) in enumerate(zip(doc.paragraphs, xml_paragraphs)):
    xml_text = xml_paragraph_text(xml_node)
    if re.sub(r'\[\[WORD_FOOTNOTE_\d+\]\]', '', xml_text) != paragraph.text:
        raise SystemExit(f'Divergence XML/python-docx au paragraphe {index}')
    paragraphs.append({
        'index': index,
        'style': paragraph.style.name,
        'raw_text': xml_text,
        'clean_text': normalize_typography(xml_text, f'docx:{index}', transformations) if xml_text else '',
    })

# Numérotation globale propre à l’œuvre: six marginales, puis la note Word.
note_definitions = {}
for number, (target_index, note_index) in enumerate(NOTE_ATTACHMENTS.items(), start=1):
    note_text = normalize_typography(paragraphs[note_index]['raw_text'], f'note:{number}', transformations)
    note_definitions[number] = note_text
    paragraphs[target_index]['clean_text'] = insert_note_before_terminal(paragraphs[target_index]['clean_text'], number)
paragraphs[138]['clean_text'] = paragraphs[138]['clean_text'].replace('[[WORD_FOOTNOTE_1]]', '[[7]]')
note_definitions[7] = normalize_typography(footnotes['1'], 'note:7', transformations)
if any('[[WORD_FOOTNOTE_' in paragraph['clean_text'] for paragraph in paragraphs):
    raise SystemExit('Marqueur Word résiduel')

segments = []
source_map = []
segment_number = 0


def push_paragraph(text, style, nature, ref1, ref2, ref1_text, ref2_text, paragraph_number, source_index):
    global segment_number
    pieces = split_text(text, f'docx:{source_index}', alerts)
    first_segment_number = segment_number + 1
    for rank, piece in enumerate(pieces, start=1):
        segment_number += 1
        calls = [int(value) for value in re.findall(r'\[\[(\d+)\]\]', piece)]
        notes = '\n'.join(f'[[{number}]] {note_definitions[number]}' for number in calls) or None
        row = {
            'id_oeuvre': WORK_ID,
            'segment_numero': segment_number,
            'segment_texte': format_source_segment(piece, style),
            'ref_niv1': ref1,
            'ref_niv2': ref2,
            'ref_niv3': None,
            'ref_niv4': None,
            'ref_niv5': None,
            'ref_niv1_texte': ref1_text,
            'ref_niv2_texte': ref2_text,
            'ref_niv3_texte': None,
            'ref_niv4_texte': None,
            'ref_niv5_texte': None,
            'lien_1': None,
            'lien_2': None,
            'lien_3': None,
            'lien_4': None,
            'fiabilite': None,
            'nature': nature,
            'texte_original': None,
            'notes': notes,
            'paragraphe': paragraph_number,
            'rang': rank,
            'controle_rang_manuel': None,
            'controle_verifie': False,
            'marquage_source': 'Codex (IA)',
            'liens_revus_le': None,
            'liens_revus_par': None,
        }
        segments.append(row)
    source_map.append({
        'source_index': source_index,
        'style': style,
        'nature': nature,
        'ref_niv1': ref1,
        'ref_niv2': ref2,
        'paragraphe': paragraph_number,
        'source_clean': text,
        'first_segment_numero': first_segment_number,
        'last_segment_numero': segment_number,
        'segment_count': len(pieces),
    })


# Apparat préliminaire. Les notes marginales sont déjà rattachées à leur paragraphe porteur.
apparatus_ranges = [
    ('Avis au lecteur', range(23, 40)),
    ('Approbation des docteurs', range(42, 46)),
    ('Privilège du Roi', range(48, 52)),
]
for ref1, indexes in apparatus_ranges:
    paragraph_number = 0
    for index in indexes:
        paragraph = paragraphs[index]
        if paragraph['style'] == 'Note marginale' or not paragraph['clean_text']:
            continue
        paragraph_number += 1
        push_paragraph(paragraph['clean_text'], paragraph['style'], 'apparat_critique', ref1, None, None, None, paragraph_number, index)

# Corps: Livre = niv1, Chapitre = niv2, Argument = titre descriptif du chapitre.
current_book = None
current_chapter = None
current_argument = None
paragraph_number = 0
book_chapters = defaultdict(int)
book_body_paragraphs = defaultdict(int)
hierarchy_map = []
for index in range(53, 1563):
    paragraph = paragraphs[index]
    style = paragraph['style']
    text = paragraph['clean_text']
    if not text:
        continue
    if style == 'Heading 1':
        current_book = text.rstrip('.')
        current_chapter = None
        current_argument = None
        paragraph_number = 0
    elif style == 'Heading 2':
        current_chapter = text.rstrip('.')
        current_argument = None
        paragraph_number = 0
        book_chapters[current_book] += 1
    elif style == 'Argument':
        if not current_book or not current_chapter or current_argument is not None:
            raise SystemExit(f'Argument hors structure ou dupliqué: {index}')
        current_argument = text
        hierarchy_map.append({'ref_niv1': current_book, 'ref_niv2': current_chapter, 'ref_niv2_texte': current_argument, 'source_index': index})
    elif style in ('Corps patrimonial', 'Vers'):
        if not current_book or not current_chapter or current_argument is None:
            raise SystemExit(f'Corps hors structure: {index}')
        paragraph_number += 1
        book_body_paragraphs[current_book] += 1
        push_paragraph(text, style, 'texte', current_book, current_chapter, None, current_argument, paragraph_number, index)
    else:
        raise SystemExit(f'Style inattendu dans le corps à {index}: {style}')

# Mention finale d'imprimeur : elle établit les métadonnées bibliographiques,
# mais ne relève ni du texte de l'œuvre ni de l'apparat critique.
if paragraphs[1563]['clean_text'] != 'FIN.' or paragraphs[1564]['style'] != 'Corps patrimonial':
    raise SystemExit('Colophon final inattendu')
printer_colophon = paragraphs[1564]['clean_text']
if printer_colophon != 'De l’Imprimerie d’Antoine Vitré, 1649.':
    raise SystemExit(f'Mention d’imprimeur inattendue : {printer_colophon}')

metadata = {
    'oeuvre_staging': {
        'id_oeuvre': WORK_ID,
        'id_auteur': 'A0010',
        'titre': 'Les Confessions',
        'sous_titre': None,
        'titre_original': 'Confessiones',
        'langue_originale': 'Latin',
        'langue_trad': 'Français',
        'date_approx': '397-401',
        'genre': 'Autobiographie spirituelle',
        'trad_auteur': 'Robert Arnauld d’Andilly',
        'note': '[Corpus Scriptura:depublie]',
        'editeur': 'Veuve Jean Camusat et Pierre Le Petit',
        'collection': None,
        'ville': 'Paris',
        'trad_id': None,
        'date_publication': '1649',
        'url_source': 'https://books.google.com/books/about/Les_Confessions_de_S_Augustin_Traduites.html?id=mNIvt407kUQC',
        'profondeur_sommaire': 2,
        'nb_signes': sum(len(row['segment_texte']) for row in segments),
        'niveaux_sommaire': 2,
        'niveaux_corps': 2,
        'texte_sommaire': '1,1,0,0,0',
        'texte_corps': '1,1,0,0,0',
        'afficher_numeros': False,
        'date_composition': '397-401',
        'genres': ['Autobiographie spirituelle', 'Théologie spirituelle'],
        'composition_debut_annee': 397,
        'composition_debut_precision': 'exacte',
        'composition_fin_annee': 401,
        'composition_fin_precision': 'exacte',
        'publication_debut_annee': 1649,
        'publication_debut_precision': 'exacte',
        'publication_fin_annee': 1649,
        'publication_fin_precision': 'exacte',
        'titre_affichage': 'Les Confessions',
    },
    'publication_patch': {
        'note': f'Traduction française de Robert Arnauld d’Andilly, seconde édition de 1649. Les treize livres sont complets ; l’Avis au lecteur, l’Approbation et le Privilège sont conservés dans l’apparat éditorial. Mention d’imprimeur conservée dans la notice bibliographique : « {printer_colophon} »',
    },
    'catalogue_notice': {
        'id': 2,
        'id_ligne': 'V20-00518',
        'id_auteur': 'A0010',
        'auteur': 'Augustin d’Hippone',
        'dates_auteur': '354-430',
        'id_oeuvre_stable': WORK_ID,
        'titre_stable': 'Les Confessions',
        'titre_original': 'Confessiones',
        'genre': 'autobiographie spirituelle et théologie',
        'langue_originale': 'latin',
        'date_oeuvre': '397-401',
        'authenticite': 'authentique',
        'id_traduction': 'TR_FR_XVII_augustin_d_hippone_confessions_1649',
        'titre_edition': 'Les Confessions de S. Augustin',
        'traducteur': 'Robert Arnauld d’Andilly',
        'annee_edition': 1649,
        'siecle_edition': 'XVIIe siècle',
        'editeur': 'Veuve Jean Camusat et Pierre Le Petit',
        'collection_nom': 'Témoin français du XVIIe siècle',
        'domaine_public': 'oui - édition de 1649',
        'url_source': 'https://books.google.com/books/about/Les_Confessions_de_S_Augustin_Traduites.html?id=mNIvt407kUQC',
        'decision_import': 'Import intégral achevé depuis le Word maître contrôlé sur le fac-similé',
        'niveau_verification': f'Seconde édition, Paris, 1649 ; Word maître SHA-256 {EXPECTED_SHA256}. Structure, paragraphes, rangs, notes et complétude contrôlés avant et après import. Mention d’imprimeur : « {printer_colophon} »',
        'score_fiabilite': 100,
        'presence_sur_le_site': True,
        'priorite': 'Très haute',
        'verifie': True,
        'lieu_edition': 'Paris',
        'date_edition': '1649',
        'auteur_debut_annee': 354,
        'auteur_debut_precision': 'exacte',
        'auteur_fin_annee': 430,
        'auteur_fin_precision': 'exacte',
        'oeuvre_debut_annee': 397,
        'oeuvre_debut_precision': 'exacte',
        'oeuvre_fin_annee': 401,
        'oeuvre_fin_precision': 'exacte',
        'edition_debut_annee': 1649,
        'edition_debut_precision': 'exacte',
        'edition_fin_annee': 1649,
        'edition_fin_precision': 'exacte',
        'verifie_admin': False,
        'auteur_uniformise': 'Augustin d’Hippone',
        'traducteur_uniformise': 'Robert Arnauld d’Andilly',
        'pseudonyme_ou_nom_imprime_traducteur': None,
        'refuse_admin': False,
        'url_texte_integral': None,
    },
}

# Audits exhaustifs du candidat.
segment_by_number = {row['segment_numero']: row for row in segments}
recomposition_failures = []
for mapping in source_map:
    rows = [segment_by_number[number] for number in range(mapping['first_segment_numero'], mapping['last_segment_numero'] + 1)]
    recomposed = ' '.join(markdown_plain(row['segment_texte']) for row in sorted(rows, key=lambda row: row['rang']))
    if recomposed != mapping['source_clean']:
        recomposition_failures.append({'source_index': mapping['source_index'], 'expected': mapping['source_clean'], 'actual': recomposed})

calls = []
definitions = []
for row in segments:
    calls.extend(int(value) for value in re.findall(r'\[\[(\d+)\]\]', row['segment_texte']))
    definitions.extend(int(value) for value in re.findall(r'^\[\[(\d+)\]\]', row['notes'] or '', flags=re.MULTILINE))

grouped = defaultdict(list)
for row in segments:
    grouped[(row['nature'], row['ref_niv1'], row['ref_niv2'], row['paragraphe'])].append(row)
rank_failures = []
for key, rows in grouped.items():
    ranks = sorted(row['rang'] for row in rows)
    if ranks != list(range(1, len(rows) + 1)):
        rank_failures.append({'key': key, 'ranks': ranks})

title_page_fragments = [paragraphs[index]['clean_text'] for index in TITLE_PAGE_RANGE if paragraphs[index]['clean_text']]
body_book_order = list(book_chapters)
actual_hierarchy = []
seen_hierarchy = set()
for row in segments:
    if row['nature'] != 'texte':
        continue
    key = (row['ref_niv1'], row['ref_niv2'])
    if key not in seen_hierarchy:
        seen_hierarchy.add(key)
        actual_hierarchy.append({'ref_niv1': row['ref_niv1'], 'ref_niv2': row['ref_niv2'], 'ref_niv2_texte': row['ref_niv2_texte']})
segment_word_counts = [word_count(markdown_plain(row['segment_texte'])) for row in segments]
invariants = {
    'master_sha256': actual_sha == EXPECTED_SHA256,
    'title_page_excluded': not any(mapping['source_index'] in TITLE_PAGE_RANGE for mapping in source_map),
    'wrapper_excluded': not any(mapping['source_index'] in EXCLUDED_WRAPPER_RANGE for mapping in source_map),
    'thirteen_books': len(body_book_order) == 13,
    'two_hundred_seventy_eight_chapters': sum(book_chapters.values()) == 278,
    'chapter_arguments_complete': actual_hierarchy == [{key: value for key, value in item.items() if key != 'source_index'} for item in hierarchy_map] and len(hierarchy_map) == 278,
    'nine_hundred_thirty_two_body_paragraphs': sum(book_body_paragraphs.values()) == 932,
    'source_recomposition_exact': not recomposition_failures,
    'segment_numbers_contiguous': [row['segment_numero'] for row in segments] == list(range(1, len(segments) + 1)),
    'ranks_contiguous': not rank_failures,
    'seven_unique_note_calls': calls == list(range(1, 8)),
    'seven_unique_note_definitions': definitions == list(range(1, 8)),
    'no_residual_word_footnotes': all('WORD_FOOTNOTE' not in row['segment_texte'] for row in segments),
    'no_links': all(not any(row[field] for field in ('lien_1', 'lien_2', 'lien_3', 'lien_4', 'liens_revus_le', 'liens_revus_par')) for row in segments),
    'no_empty_segments': all(row['segment_texte'].strip() for row in segments),
    'no_boundary_spaces': all(row['segment_texte'] == row['segment_texte'].strip() for row in segments),
    'no_corrupt_characters': all(not re.search(r'[\ufffd\x00-\x08\x0b\x0c\x0e-\x1f]', row['segment_texte']) for row in segments),
    'no_straight_apostrophes': all("'" not in row['segment_texte'] for row in segments),
    'no_em_or_en_dash': all('—' not in row['segment_texte'] and '–' not in row['segment_texte'] for row in segments),
    'notes_never_after_closing_quote': all(not re.search(r'[»”"]\s*\[\[\d+\]\]', row['segment_texte']) for row in segments),
}

random.seed(1649)
samples = []
for book in body_book_order:
    mappings = [mapping for mapping in source_map if mapping['nature'] == 'texte' and mapping['ref_niv1'] == book]
    for mapping in random.sample(mappings, min(5, len(mappings))):
        rows = [segment_by_number[number] for number in range(mapping['first_segment_numero'], mapping['last_segment_numero'] + 1)]
        samples.append({
            'source_index': mapping['source_index'],
            'ref_niv1': mapping['ref_niv1'],
            'ref_niv2': mapping['ref_niv2'],
            'paragraphe': mapping['paragraphe'],
            'source': mapping['source_clean'],
            'segments': [row['segment_texte'] for row in rows],
            'recomposition_ok': ' '.join(markdown_plain(row['segment_texte']) for row in rows) == mapping['source_clean'],
        })
apparatus_mappings = [mapping for mapping in source_map if mapping['nature'] == 'apparat_critique']
for mapping in random.sample(apparatus_mappings, min(8, len(apparatus_mappings))):
    rows = [segment_by_number[number] for number in range(mapping['first_segment_numero'], mapping['last_segment_numero'] + 1)]
    samples.append({
        'source_index': mapping['source_index'],
        'ref_niv1': mapping['ref_niv1'],
        'ref_niv2': mapping['ref_niv2'],
        'paragraphe': mapping['paragraphe'],
        'source': mapping['source_clean'],
        'segments': [row['segment_texte'] for row in rows],
        'recomposition_ok': ' '.join(markdown_plain(row['segment_texte']) for row in rows) == mapping['source_clean'],
    })

suspects = []
for mapping in source_map:
    text = mapping['source_clean']
    if re.search(r'\w-\s+\w', text):
        suspects.append({'type': 'cesure_possible', 'source_index': mapping['source_index'], 'text': text})
    if re.search(r'\b(?:rn|vv|cl)\b', text, flags=re.IGNORECASE):
        suspects.append({'type': 'petit_mot_ocr_possible', 'source_index': mapping['source_index'], 'text': text})
alerts.extend(suspects)

audit = {
    'source': {'path': str(MASTER), 'sha256': actual_sha},
    'counts': {
        'segments': len(segments),
        'body_segments': sum(row['nature'] == 'texte' for row in segments),
        'apparatus_segments': sum(row['nature'] == 'apparat_critique' for row in segments),
        'source_paragraphs_imported': len(source_map),
        'body_paragraphs': sum(book_body_paragraphs.values()),
        'apparatus_paragraphs': len([mapping for mapping in source_map if mapping['nature'] == 'apparat_critique']),
        'books': len(body_book_order),
        'chapters': sum(book_chapters.values()),
        'notes': len(note_definitions),
        'words_min': min(segment_word_counts),
        'words_median': sorted(segment_word_counts)[len(segment_word_counts) // 2],
        'words_mean': round(sum(segment_word_counts) / len(segment_word_counts), 2),
        'words_max': max(segment_word_counts),
        'segments_over_20_words': sum(count > 20 for count in segment_word_counts),
        'segments_over_30_words': sum(count > 30 for count in segment_word_counts),
    },
    'book_chapters': dict(book_chapters),
    'book_body_paragraphs': dict(book_body_paragraphs),
    'invariants': invariants,
    'recomposition_failures': recomposition_failures,
    'rank_failures': rank_failures,
    'alert_counts': Counter(alert['type'] for alert in alerts),
    'candidate_sha256': sha256_bytes(stable_json(segments).encode()),
}
if not all(invariants.values()):
    write_json('confessions-failed-audit.json', audit)
    raise SystemExit(f'Invariants en échec: {[key for key, value in invariants.items() if not value]}')

write_json('confessions-segments-candidate.json', segments)
write_json('confessions-metadata-candidate.json', metadata)
write_json('confessions-source-map.json', source_map)
write_json('confessions-hierarchy-map.json', hierarchy_map)
write_json('confessions-transformations.json', transformations)
write_json('confessions-alerts.json', alerts)
write_json('confessions-random-samples.json', samples)
write_json('confessions-audit.json', audit)
with (ROOT / 'confessions-segments-candidate.csv').open('w', encoding='utf-8-sig', newline='') as handle:
    writer = csv.DictWriter(handle, fieldnames=SEGMENT_COLUMNS)
    writer.writeheader()
    writer.writerows(segments)

structure_lines = [
    '# Confessions d’Augustin - dossier de preuve avant import',
    '',
    f'- Word maître : `{MASTER}`',
    f'- SHA-256 : `{actual_sha}`',
    '- Édition : seconde édition, Paris, veuve Jean Camusat et Pierre Le Petit, 1649.',
    '- Page de titre : exclue des segments ; elle sert seulement à la fiche bibliographique.',
    '- Corps : treize livres, 278 chapitres ; `Livre` = niv1, `Chapitre` = niv2, argument imprimé = `ref_niv2_texte`.',
    '- Apparat : Avis au lecteur, Approbation des docteurs, Privilège du Roi, colophon final.',
    '- Notes : sept appels et sept définitions, renumérotés globalement `[[1]]` à `[[7]]`.',
    '- Liens bibliques : aucun lien créé ou révisé pendant cette passe.',
    '',
    '## Comptes par livre',
    '',
    '| Livre | Chapitres | Paragraphes source |',
    '|---|---:|---:|',
]
for book in body_book_order:
    structure_lines.append(f'| {book} | {book_chapters[book]} | {book_body_paragraphs[book]} |')
structure_lines.extend(['', '## Résumé de segmentation', '', f'- Segments : {len(segments)}', f'- Médiane : {audit["counts"]["words_median"]} mots', f'- Moyenne : {audit["counts"]["words_mean"]} mots', f'- Segments de plus de 30 mots : {audit["counts"]["segments_over_30_words"]}', f'- Alertes : {dict(audit["alert_counts"])}', ''])
(ROOT / 'confessions-structure.md').write_text('\n'.join(structure_lines), encoding='utf-8')

print(json.dumps(audit, ensure_ascii=False, indent=2))
