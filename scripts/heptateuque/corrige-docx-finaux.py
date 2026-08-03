from __future__ import annotations

import shutil
from pathlib import Path
from docx import Document


ROOT = Path(__file__).resolve().parent
BACKUP = ROOT / "audit-reprise" / "docx-avant-corrections-finales-2026-08-02"
FILES = [
    "Genese_draft_v7.docx", "Exode_draft_v3.docx", "Levitique_draft.docx",
    "Deuteronome_draft.docx", "Josue_draft.docx", "Juges_draft.docx",
]

REPLACEMENTS = {
    "remplir sa cité": "remplir sa cité ?",
    "un froid engourdi, pour : un": "un froid engourdi, pour : un froid qui engourdit.",
    "naître le Christ": "naître le Christ ?",
    "culte de latrie": "culte de latrie.",
    "par cela même qu’il est Dieu": "par cela même qu’il est Dieu.",
    "le Fils unique de Dieu": "le Fils unique de Dieu ?",
    "signification du mot Chérubin,": "signification du mot Chérubin.",
    " ==DU TABERNACLE.==": "",
    " v 1-8": "",
    "purs de toutes ces infamies,": "purs de toutes ces infamies.",
}

REFERENCES = {
    "32, 20.": "Genèse 32, 20",
    "Josué, I, 5.": "Josué 1, 5",
    "Deutéronome, I, 29, 30.": "Deutéronome 1, 29, 30",
}


def replace_suffix_in_runs(paragraph, old: str, new: str) -> bool:
    if not paragraph.text.endswith(old):
        return False
    remaining = old
    touched = []
    for run in reversed(paragraph.runs):
        if not remaining:
            break
        take = min(len(run.text), len(remaining))
        if run.text[-take:] != remaining[-take:]:
            return False
        touched.append((run, take))
        remaining = remaining[:-take]
    if remaining:
        return False
    first_run, first_take = touched[-1]
    first_run.text = first_run.text[:-first_take] + new
    for run, take in touched[:-1]:
        run.text = run.text[:-take]
    return True


def correct(path: Path) -> int:
    document = Document(path)
    changes = 0
    for paragraph in document.paragraphs:
        for old, new in REPLACEMENTS.items():
            if replace_suffix_in_runs(paragraph, old, new):
                changes += 1
                break
        if paragraph.text in REFERENCES:
            paragraph.runs[0].text = REFERENCES[paragraph.text]
            for run in paragraph.runs[1:]:
                run.text = ""
            changes += 1

    # Titre et premier intertitre de la longue récapitulation de l'Exode.
    if path.name == "Exode_draft_v3.docx":
        paragraphs = document.paragraphs
        for index, paragraph in enumerate(paragraphs):
            if paragraph.text == "Question CLXXVII":
                target = paragraphs[index + 1]
                if not target.text:
                    run = target.add_run("Du tabernacle. — But de ce travail.")
                    run.italic = True
                    changes += 1
                break

    # La signature finale est une phrase d'apparat ; le point reste hors italique.
    if path.name == "Juges_draft.docx":
        for paragraph in document.paragraphs:
            if paragraph.text.endswith("Cette traduction est l’œuvre de M. l’abbé POGNON"):
                paragraph.add_run(".")
                changes += 1
                break

    if changes:
        document.save(path)
    return changes


def main() -> None:
    BACKUP.mkdir(parents=True, exist_ok=True)
    total = 0
    for name in FILES:
        path = ROOT / name
        backup = BACKUP / name
        if not backup.exists():
            shutil.copy2(path, backup)
        count = correct(path)
        total += count
        print(f"{name}: {count} correction(s)")
    if total != 15:
        raise RuntimeError(f"Nombre de corrections inattendu : {total}/15")


if __name__ == "__main__":
    main()
