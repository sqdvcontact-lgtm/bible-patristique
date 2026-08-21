from __future__ import annotations

import json
import re
import unicodedata
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
AUDIT = ROOT / "audit" / "hexameron-2026-07-30"
DOCX = Path(r"C:\Corpus Scriptura\Sources\Basile\BKV_Homelies_sur_Hexaemeron_Auger_1827.docx")
COMPARISON = json.loads((AUDIT / "pdf-db-comparison.json").read_text(encoding="utf-8"))
OUTPUT = AUDIT / "apparatus-proposal.json"

WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿŒœÆæ]+(?:[’'][A-Za-zÀ-ÖØ-öø-ÿŒœÆæ]+)*")
HOMILY_LABELS = {
    1: "Première homélie",
    2: "Deuxième homélie",
    3: "Troisième homélie",
    4: "Quatrième homélie",
    5: "Cinquième homélie",
    6: "Sixième homélie",
    7: "Septième homélie",
    8: "Huitième homélie",
    9: "Neuvième homélie",
    10: "Dixième homélie (attribution discutée)",
}
SUMMARY_INDICES = {1: 17, 2: 54, 3: 84, 4: 115, 5: 140, 6: 172, 7: 210, 8: 233, 9: 263}


def key(value: str) -> str:
    value = value.lower().replace("’", "'").replace("œ", "oe").replace("æ", "ae")
    return "".join(c for c in unicodedata.normalize("NFD", value) if unicodedata.category(c) != "Mn")


source_forms: dict[str, Counter[str]] = defaultdict(Counter)
reverse: dict[str, list[str]] = defaultdict(list)
for homily in COMPARISON["homilies"]:
    for item in homily["graphies_anciennes"]:
        source_forms[key(item["source"])][item["source"].replace("'", "’")] += 1
for source_key, database_key in COMPARISON["equivalences_apprises"].items():
    reverse[database_key].append(source_key)


def source_form(word: str) -> str | None:
    database_key = key(word)
    candidates = reverse.get(database_key, [])
    if not candidates:
        return None
    source_key = max(candidates, key=lambda candidate: sum(source_forms[candidate].values()))
    forms = {word}
    for _ in range(2):
        for value in list(forms):
            forms.add(re.sub(r"aient$", "oient", value, flags=re.IGNORECASE))
            forms.add(re.sub(r"ait$", "oit", value, flags=re.IGNORECASE))
            forms.add(re.sub(r"ais$", "ois", value, flags=re.IGNORECASE))
            forms.add(re.sub(r"ements$", "emens", value, flags=re.IGNORECASE))
            forms.add(re.sub(r"ants$", "ans", value, flags=re.IGNORECASE))
            forms.add(re.sub(r"temps$", "tems", value, flags=re.IGNORECASE))
            forms.add(value.replace("conna", "conno").replace("Conna", "Conno"))
            forms.add(value.replace("faibl", "foibl").replace("Faibl", "Foibl"))
            forms.add(value.replace("para", "paro").replace("Para", "Paro"))
            forms.add(value.replace("paiement", "payement").replace("Paiement", "Payement"))
    matches = [value for value in forms if key(value) == source_key]
    return min(matches, key=lambda value: (value == word, len(value))) if matches else None


def restore_source_spelling(text: str) -> tuple[str, list[dict[str, str]]]:
    replacements = []
    pieces = []
    cursor = 0
    for match in WORD_RE.finditer(text):
        pieces.append(text[cursor:match.start()])
        before = match.group(0)
        after = source_form(before) or before
        pieces.append(after)
        if before != after:
            replacements.append({"before": before, "after": after})
        cursor = match.end()
    pieces.append(text[cursor:])
    return "".join(pieces), replacements


def normalize_typography(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace("'", "’")
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"([,.;:!?])(?=[A-Za-zÀ-ÖØ-öø-ÿŒœÆæ])", r"\1 ", text)
    return text


MANUAL_CORRECTIONS = {
    "Sommaire général": [
        ("longtemps encore", "long-temps encore"),
        ("on donnait", "on donnoit"),
        ("embellissaient son ouvrage", "embellissoient son ouvrage"),
        ("saint Basile écrivait", "saint Basile écrivoit"),
        ("auraient été", "auroient été"),
        ("il joignait", "il joignoit"),
        ("ses oeuvres", "ses œuvres"),
        ("les âmes religieuses", "les ames religieuses"),
        ("Je ne savais", "Je ne savois"),
        ("j’en avais appris", "j’en avois appris"),
        ("les talents", "les talens"),
        ("de rue marquer", "de me marquer"),
        ("voudront traiter", "voudroit traiter"),
        ("le peuple même connaissait", "le peuple même connoissoit"),
        ("qui le doit jamais finir", "qui ne doit jamais finir"),
        ("les éléments de l’eau", "les élémens de l’eau"),
        ("Les eaux qui couvraient", "Les eaux qui couvroient"),
        ("qui ôtait au aronde visible", "qui ôtoit au monde visible"),
        ("qui le laissait dans", "qui le laissoit dans"),
        ("Ou ne peut dire", "On ne peut dire"),
        ("retour de la lainière", "retour de la lumière"),
        ("dans lequel n été transporté", "dans lequel a été transporté"),
        ("firmament, couture sur", "firmament, comme sur"),
        ("en sou lieu", "en son lieu"),
        ("Création de «deux", "Création de deux"),
        ("Celte raison", "Cette raison"),
        ("n’en trouverait pas", "n’en trouveroit pas"),
        ("il suffirait", "il suffiroit"),
        ("n’offre. rien", "n’offre rien"),
        ("produisait d’elle-même", "produisoit d’elle-même"),
        ("sans oeuf", "sans œuf"),
    ],
    "Première homélie": [
        ("Moise", "Moïse"),
        ("qui attribuaient le monde", "qui attribuoient le monde"),
        ("toutes ara premières", "toutes aux premières"),
        ("ex-posé", "exposé"),
        ("les éléments de l’eau", "les élémens de l’eau"),
        ("on chercherait", "on chercheroit"),
        ("on examinerait", "on examineroit"),
    ],
    "Deuxième homélie": [
        ("Moise", "Moïse"),
        ("plusieurs retendaient", "plusieurs entendoient"),
        ("mise en oeuvre", "mise en œuvre"),
        ("qu’ils prétendaient", "qu’ils prétendoient"),
        ("il les refile", "il les réfute"),
        ("ne saurait être", "ne sauroit être"),
        ("couvraient la face", "couvroient la face"),
        ("c’était, selon eux", "c’étoit, selon eux"),
        ("l’esprit de Dieu échauffait", "l’esprit de Dieu échauffoit"),
    ],
    "Troisième homélie": [
        ("second tour", "second jour"),
        ("qui pensaient", "qui pensoient"),
        ("nous paroît", "nous paroît"),
        ("Des écrivains expliquaient", "Des écrivains expliquoient"),
    ],
    "Quatrième homélie": [
        ("Il explique continent", "Il explique comment"),
        ("dans un meure lieu", "dans un même lieu"),
        ("les eaux ont revu", "les eaux ont reçu"),
        ("la terre, connue la chaleur", "la terre, comme la chaleur"),
        ("dont les éléments se rapprochent", "dont les élémens se rapprochent"),
    ],
    "Cinquième homélie": [
        ("Elle égale par", "Elle égale par"),
    ],
    "Sixième homélie": [
        ("ce qu’elle étoit destinée", "ce qu’elle étoit destinée"),
        ("qu’ou peut voir", "qu’on peut voir"),
        ("contre cens qui", "contre ceux qui"),
        ("astres, prétendaient", "astres, prétendoient"),
        ("Il étabit", "Il établit"),
        ("ces deus astres oui", "ces deux astres qui"),
        ("qu’a cause", "qu’à cause"),
    ],
    "Septième homélie": [
        ("propres â chaque", "propres à chaque"),
    ],
    "Huitième homélie": [
        ("comme s’il allait parler", "comme s’il alloit parler"),
        ("une âme vivante", "une ame vivante"),
        ("qui donnaient une âme", "qui donnoient une ame"),
        ("une âme qui gouverne", "une ame qui gouverne"),
        ("de l’âme humaine", "de l’ame humaine"),
        ("Eu effet", "En effet"),
        ("qu’il avoit", "qu’il avoit"),
        ("qu’U va parler", "qu’il va parler"),
        ("des reptiles animés…", "des reptiles animés…"),
        ("si l’un n’en trouve", "si l’on n’en trouve"),
        ("insectes volans", "insectes volans"),
        ("termine son homélie eu demandant", "termine son homélie en demandant"),
    ],
    "Neuvième homélie": [
        ("création, ou ont été", "création, où ont été"),
        ("Pans cette homélie", "Dans cette homélie"),
        ("soir, ait il sera", "soir, où il sera"),
        ("cherchaient des sens", "cherchoient des sens"),
        ("demis les différentes", "dans les différentes"),
        ("l’âme des bêtes", "l’ame des bêtes"),
        ("de l’âme humaine", "de l’ame humaine"),
        ("plus en détail", "plus en détail"),
    ],
}


def correct_transcription(text: str, group: str) -> tuple[str, list[dict[str, str]]]:
    corrections = []
    for before, after in MANUAL_CORRECTIONS.get(group, []):
        if before == after:
            continue
        if before not in text:
            continue
        text = text.replace(before, after)
        corrections.append({"before": before, "after": after})
    return text, corrections


doc = Document(DOCX)
rows = []

# Le sommaire général comporte un paragraphe de présentation, puis sept paragraphes
# sous la rubrique éditoriale « Système de St. Basile… ».
general_indices = [6, 8, 9, 10, 11, 12, 13, 14]
for paragraph, docx_index in enumerate(general_indices, start=1):
    before = normalize_typography(doc.paragraphs[docx_index].text)
    after, replacements = restore_source_spelling(before)
    after, corrections = correct_transcription(after, "Sommaire général")
    rows.append({
        "group": "Sommaire général",
        "ref_niv1": "Sommaire général",
        "ref_niv1_texte": None,
        "ref_niv2": "Système de saint Basile sur la création du monde" if docx_index >= 8 else None,
        "ref_niv2_texte": None,
        "paragraphe": paragraph,
        "rang": 1,
        "source": f"DOCX paragraphe {docx_index}; fac-similé PDF 381-384",
        "segment_texte": after,
        "replacements": replacements,
        "corrections": corrections,
    })

for number, docx_index in SUMMARY_INDICES.items():
    before = normalize_typography(doc.paragraphs[docx_index].text)
    after, replacements = restore_source_spelling(before)
    after, corrections = correct_transcription(after, HOMILY_LABELS[number])
    rows.append({
        "group": HOMILY_LABELS[number],
        "ref_niv1": HOMILY_LABELS[number],
        "ref_niv1_texte": "Sommaire",
        "ref_niv2": None,
        "ref_niv2_texte": None,
        "paragraphe": 1,
        "rang": 1,
        "source": f"DOCX paragraphe {docx_index}; fac-similé contrôlé",
        "segment_texte": after,
        "replacements": replacements,
        "corrections": corrections,
    })

tenth_source = """
On trouve dans les œuvres de saint Basile, et dans celles de saint Grégoire de Nysse son frère, deux homélies sur la création de l’homme, qui ne sont certainement ni de l’un ni de l’autre. Saint Grégoire de Nysse, dans un traité sur la création de l’homme, dit en termes formels que Basile son frère a laissé imparfait l’hexaëméron, et qu’il a composé son traité pour y suppléer. Ce traité, où il y a de belles choses, annonce qu’il n’est pas l’auteur des deux homélies. L’autorité de saint Grégoire de Nysse suffiroit seule pour convaincre que saint Basile n’en est pas non plus l’auteur : une ancienne version latine de l’hexaëméron, où les deux homélies ne sont pas traduites, le témoignage tacite de saint Ambroise, qui a mis en latin l’hexaëméron grec, et qui traite de lui-même la formation de l’homme, sans rien prendre des deux homélies, ces deux nouvelles autorités portent la chose au dernier degré de démonstration. J’ai lu cependant les deux homélies ; j’y ai trouvé bien des choses opposées à la manière de St. Basile, et peu dignes de ce grand orateur ; mais j’y en ai trouvé aussi beaucoup qu’il n’auroit pas désavouées. En général, elles offrent plus de beautés oratoires que le traité de saint Grégoire de Nysse. J’ai donc supprimé tout ce qui m’a paru être peu intéressant, ou ralentir la marche du discours, et des deux homélies je n’en ai fait qu’une seule telle que je la publie aujourd’hui en françois. En voici le sommaire.

L’orateur se met à la place de saint Basile ; il annonce qu’il vient s’acquitter d’une ancienne dette dont la maladie lui a fait différer le payement. Il se plaint que l’homme ne s’étudie pas lui-même, et qu’il ne cherche pas à admirer l’Ouvrier suprême en admirant les merveilles que lui offre sa propre existence. Il explique ces paroles : Faisons l’homme à notre image et à notre ressemblance. Dieu semble délibérer avant de créer l’homme, ce qui annonce l’importance de cet être. Nous avons été faits à l’image de Dieu. Qu’on écarte toute idée de figure corporelle, qui est incompatible avec l’idée de Dieu. Comment donc avons-nous été faits à l’image de Dieu ? les paroles suivantes l’expliquent : Faisons l’homme à notre image… et qu’il commande aux poissons… C’est par la raison, c’est par les lumières de son intelligence, et non par les forces de son corps, que l’homme commande aux animaux. Il est né pour commander, qu’il prenne garde de s’asservir aux passions. On montre comment l’homme commande aux poissons, aux bêtes sauvages, aux oiseaux du ciel ; tout ce morceau renferme de grandes beautés. On explique avec autant de solidité que de subtilité en quoi diffèrent ces deux expressions : À notre image et à notre ressemblance. Dieu prit du limon de la terre et forma l’homme. Dieu nous travaille de sa propre main, ayons une grande idée de nous-mêmes ; il nous forme du limon de la terre, n’ayons que des sentimens modestes. Ces deux idées sont développées avec éloquence. Le mot forma annonce un certain art dont use l’Ouvrier suprême en créant l’homme. L’homme est vraiment un petit monde qui offre au contemplateur attentif un nombre infini de merveilles. L’orateur s’arrête surtout à la stature droite du corps humain, et à la beauté de l’œil dont il donne une description fort étendue qui termine l’homélie.
"""
for paragraph, raw_text in enumerate(tenth_source.strip().split("\n\n"), start=1):
    text = normalize_typography(raw_text)
    rows.append({
        "group": HOMILY_LABELS[10],
        "ref_niv1": HOMILY_LABELS[10],
        "ref_niv1_texte": "Sommaire",
        "ref_niv2": None,
        "ref_niv2_texte": None,
        "paragraphe": paragraph,
        "rang": 1,
        "source": "Fac-similé PDF 556-557, transcription contrôlée",
        "segment_texte": text,
        "replacements": [],
        "corrections": [],
    })

payload = {
    "source_docx": str(DOCX),
    "segments": len(rows),
    "groups": dict(Counter(row["group"] for row in rows)),
    "spelling_replacements": sum(len(row["replacements"]) for row in rows),
    "rows": rows,
}
OUTPUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
print(json.dumps({key: payload[key] for key in ("segments", "groups", "spelling_replacements")}, ensure_ascii=False, indent=2))
print(OUTPUT)
